import io
import re
import docx
from docx.shared import Pt, RGBColor

def clean_ocr_text(text: str, target_lang: str = "auto") -> str:
    text = str(text or "").strip()
    if not text:
        return ""
        
    refusals = [
        "죄송합니다", "번역해", "번역할", "해드릴 수", "없습니다", "다른 문장",
        "sorry", "cannot translate", "cannot transcribe", "unable to", "don't have",
        "전화번호가 없습니다"
    ]
    if any(ref in text.lower() for ref in refusals):
        return ""
        
    # Strip trailing translation notes
    text = re.sub(r"\s*,?\s*which translates to.*$", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*,?\s*meaning\s+.*$", "", text, flags=re.IGNORECASE)
    
    # Strip common prefixes
    prefixes = [
        r"^the text in the image is:?\s*",
        r"^the text in the image translates to:?\s*",
        r"^the transcribed text is:?\s*",
        r"^transcribed text:?\s*",
        r"^this translates to:?\s*",
        r"^in the image, the text is:?\s*",
        r"^the image displays the text:?\s*",
        r"^the bubble text is:?\s*",
        r"^text in the image:?\s*",
        r"^image text:?\s*",
    ]
    for pref in prefixes:
        text = re.sub(pref, "", text, flags=re.IGNORECASE)
        
    text = text.strip()
    
    symbol = ""
    symbol_match = re.match(r"^(\"\"|\(\)|\[\]|OT|ST|SFX|<>|::):\s*(.*)$", text)
    if symbol_match:
        symbol, rest = symbol_match.groups()
        text = rest.strip()
        
    # Strip quotes
    text = re.sub(r'^["\'“]+', '', text)
    te    # Filter out Chinese sentence hallucinations ONLY if target language is explicitly Korean ('ko')
    if target_lang == "ko":
        has_chinese = bool(re.search(r"[\u4e00-\u9fff]", text))
        has_korean = bool(re.search(r"[\uac00-\ud7a3\u3130-\u318f]", text))
        if has_chinese and not has_korean:
            return ""
            
    if symbol and text:
        return f"{symbol}: {text}"
    return text


def format_bubble_text(text: str, settings: dict = None) -> str:
    target_lang = settings.get("target_lang", "auto") if settings else "auto"
    text = clean_ocr_text(text, target_lang)
    text = str(text or "").strip()
    if not text:
        return ""
    if not settings:
        return text

    pattern = r"^(\"\"|\(\)|\[\]|OT|ST|SFX|<>|::):\s*(.*)$"
    m = re.match(pattern, text)
    if m:
        prefix, content = m.groups()
    else:
        prefix, content = "", text

    if settings.get("remove_legends"):
        return content.strip()

    mapping = {
        '""': settings.get("legend_speech", '""'),
        '::': settings.get("legend_shouting", '::'),
        'ST': settings.get("legend_small", 'ST'),
        '()': settings.get("legend_thinking", '()'),
        '[]': settings.get("legend_box", '[]'),
        '<>': settings.get("legend_system", '<>'),
        'OT': settings.get("legend_outer", 'OT'),
        'SFX': settings.get("legend_sfx", 'SFX'),
    }

    custom_prefix = mapping.get(prefix, prefix)
    if custom_prefix:
        return f"{custom_prefix}: {content.strip()}"
    return content.strip()


def create_docx(folder_name: str, pages_data: list, settings: dict = None) -> io.BytesIO:
    doc = docx.Document()
    doc.add_heading(f"Translation - {folder_name}", 0)
    
    add_sp = settings.get("add_spaces", 1) if settings else 1
    
    # Guarantee document is never empty even if pages_data is empty
    if not pages_data:
        p_head = doc.add_paragraph()
        p_head.add_run("------ Page 1 ------").bold = True
        doc.add_paragraph("[لم يتم العثور على صور أو نصوص في هذا المجلد / No images or text found]")
    
    for p in pages_data:
        page_num = p.get("page_num", 0)
        bubbles = p.get("bubbles") or []
        if not bubbles:
            texts = p.get("texts", [])
            bubbles = [{"bubble_num": i + 1, "text": text} for i, text in enumerate(texts)]

        p_head = doc.add_paragraph()
        runner = p_head.add_run(f"------ Page {page_num} ------")
        runner.bold = True
        runner.font.size = Pt(14)
        runner.font.color.rgb = RGBColor(0, 102, 204)
        
        if add_sp:
            doc.add_paragraph("")
            
        written_count = 0
        for bubble in bubbles:
            text = str(bubble.get("text", "")).strip()
            formatted = format_bubble_text(text, settings)
            if not formatted:
                continue
            p_text = doc.add_paragraph()
            r = p_text.add_run(formatted)
            r.font.size = Pt(11)
            written_count += 1
            if add_sp:
                doc.add_paragraph("")
                
        if written_count == 0:
            p_none = doc.add_paragraph()
            r_none = p_none.add_run("[لا يوجد نص في هذه الصفحة / No text detected on this page]")
            r_none.font.italic = True
            r_none.font.color.rgb = RGBColor(128, 128, 128)
            if add_sp:
                doc.add_paragraph("")

        if add_sp:
            doc.add_paragraph("")
        
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def create_plain_text(folder_name: str, pages_data: list, settings: dict = None) -> io.BytesIO:
    lines: list[str] = [f"Extraction - {folder_name}", ""]
    add_sp = settings.get("add_spaces", 1) if settings else 1

    if not pages_data:
        lines.extend(["------ Page 1 ------", "[لم يتم العثور على صور أو نصوص في هذا المجلد / No images or text found]\n"])

    for p in pages_data:
        page_num = p.get("page_num", 0)
        bubbles = p.get("bubbles") or []
        if not bubbles:
            texts = p.get("texts", [])
            bubbles = [{"bubble_num": i + 1, "text": text} for i, text in enumerate(texts)]

        lines.append(f"------ Page {page_num} ------\n")
        written_count = 0
        for bubble in bubbles:
            text = str(bubble.get("text", "")).strip()
            formatted = format_bubble_text(text, settings)
            if not formatted:
                continue
            lines.append(formatted)
            written_count += 1
            if add_sp:
                lines.append("")

        if written_count == 0:
            lines.append("[لا يوجد نص في هذه الصفحة / No text detected on this page]")
            if add_sp:
                lines.append("")

        if add_sp:
            lines.append("")

    buf = io.BytesIO(("\n".join(lines).rstrip() + "\n").encode("utf-8"))
    buf.seek(0)
    return buf
